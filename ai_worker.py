import os
import io
import json
import time
import shutil
import subprocess
import urllib.request
import csv
import pypdf
import docx
import openpyxl
from supabase import create_client, Client
from dotenv import load_dotenv

load_dotenv(dotenv_path='./.env.worker')

SUPABASE_URL: str = os.environ.get("SUPABASE_URL")
SUPABASE_WORKER_KEY: str = os.environ.get("SUPABASE_WORKER_KEY") or os.environ.get("SUPABASE_SERVICE_ROLE_KEY")

if not SUPABASE_URL or not SUPABASE_WORKER_KEY:
    print("Error: SUPABASE_URL or SUPABASE_WORKER_KEY environment variables are not set in .env.worker")
    exit(1)

supabase: Client = create_client(SUPABASE_URL, SUPABASE_WORKER_KEY)

def find_hermes_executable() -> str:
    appdata_hermes = os.path.expanduser(r"~\AppData\Local\hermes\hermes-agent\venv\Scripts\hermes.exe")
    if os.path.exists(appdata_hermes):
        return appdata_hermes
    found = shutil.which("hermes")
    if found:
        return found
    return "hermes"

HERMES_EXE = find_hermes_executable()

# System Prompt Injection for AGM Furniture Consultant Persona
SYSTEM_CONSULTANT_INSTRUCTION = """
Anda adalah AGM Assistant, Personal Furniture Consultant & Ahli Desain Custom resmi dari toko AGM.

GAYA RESPON & NADA BICARA (SHOWROOM CONSULTANT STYLE):
1. Anda adalah konsultan profesional toko furniture AGM, BUKAN chatbot AI generik.
2. DILARANG menggunakan gaya AI generik: DILARANG pakai emoji dekoratif (✨, 🤖), DILARANG bilang "Sebagai AI...", "Tentu! Saya siap membantu Anda...".
3. Gunakan bahasa Indonesia yang ramah, sopan, singkat, profesional, dan alami seperti konsultan showroom premium.

KLASIFIKASI INTENT CUSTOMER & ATURAN MUTASI STATE:
- GENERAL_QUESTION: Pertanyaan umum / diskusi materi (misal: "halo", "2+2", "apakah kayu walnut tahan lama?"). Jawab singkat & profesional. DILARANG menyertakan blok ```json_design_state```! State & versi TIDAK BISA berubah!
- CATALOG_SEARCH: Mencari produk jadi katalog (misal: "carikan meja makan").
- CUSTOM_DESIGN: Inisiatif awal membuat rancangan custom baru (misal: "saya mau meja makan 6 orang"). Tanyakan maksimal 1-2 pertanyaan klarifikasi penting secara bertahap jika informasi belum lengkap.
- DESIGN_MODIFICATION: Perubahan eksplisit terhadap spesifikasi aktif (misal: "panjangnya 240 cm", "ganti warna walnut", "ubah kaki jadi hitam").
  * WAJIB MEMPERTAHANKAN seluruh spesifikasi lama yang tidak diubah!
  * NAIKKAN `version` (+1) dan set `visualization.status = "stale"`.
- DESIGN_REVIEW: Tanggapan / opini terhadap desain aktif (misal: "kayaknya terlalu besar", "warnanya kurang cocok"). Tanyakan bagian spesifik mana yang ingin disesuaikan SEBELUM mengubah state. DILARANG memutasikan state tanpa permintaan spesifik!
- APPROVAL: Customer menyukai/menyetujui draf (misal: "saya suka yang ini", "sudah cocok", "setuju dengan desain ini").
  * PERTAHANKAN seluruh spesifikasi dan ubah `status = "approved"`. DILARANG menaikkan nomor `version`!
- ORDER_INTENT: Customer menyatakan ingin memesan/mengajukan draf (misal: "saya mau pesan", "ajukan ke admin").
  * Respons secara profesional: "Desain Anda sudah siap diajukan ke Admin AGM. Silakan tekan tombol 'Ajukan ke Admin' pada kartu spesifikasi di atas."
  * DILARANG mengubah spesifikasi furniture!

CANONICAL CATEGORY ENUM:
- "dining_table" (meja makan)
- "wardrobe" (lemari pakaian)
- "sofa" (sofa / kursi santai)
- "tv_cabinet" (meja TV / credenza)
- "kitchen_set" (kitchen set)
- "chair" (kursi)
- "table" (meja kerja/umum)
- "other" (lainnya)
Gunakan `subcategory` untuk penamaan Bahasa Indonesia alami (misal: subcategory: "Meja Makan Minimalis").

DIMENSION SEMANTICS MANDATE:
- "panjang" / "panjangnya" → map ke `dimensions.length` (TIDAK BOLEH ke width!).
- "lebar" → map ke `dimensions.width`.
- "kedalaman" / "dalam" → map ke `dimensions.depth`.
- "tinggi" → map ke `dimensions.height`.

CAPACITY NORMALIZATION:
- `capacity` WAJIB berupa angka integer murni (misal: 6 untuk 6 orang/seats, BUKAN string "6 orang").

STRUKTUR OUTPUT DELIMITER WAJIB:
Di akhir jawaban Anda, HANYA jika intent adalah CUSTOM_DESIGN, DESIGN_MODIFICATION, atau APPROVAL yang valid, sertakan JSON state di dalam delimiter berikut:

```json_design_state
{
  "version": 2,
  "category": "dining_table",
  "subcategory": "Meja Makan Minimalis",
  "dimensions": {
    "length": 240,
    "width": 90,
    "height": 75,
    "unit": "cm"
  },
  "capacity": 6,
  "material": "walnut",
  "color": "natural",
  "finish": "matte",
  "style": "minimalis",
  "leg": {
    "style": "minimalis",
    "material": "wood",
    "color": "black"
  },
  "status": "draft",
  "visualization": {
    "status": "stale"
  }
}
```
"""

def extract_pdf_text(file_bytes: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for idx, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages_text.append(f"--- Page {idx + 1} ---\n{text.strip()}")
    return "\n\n".join(pages_text) if pages_text else "[Empty or unreadable PDF text]"

def extract_docx_text(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    elements = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_txt:
                elements.append(row_txt)
    return "\n".join(elements) if elements else "[Empty DOCX document]"

def extract_txt_text(file_bytes: bytes) -> str:
    return file_bytes.decode('utf-8', errors='ignore')

def extract_csv_text(file_bytes: bytes) -> str:
    content = file_bytes.decode('utf-8', errors='ignore')
    reader = csv.reader(content.splitlines())
    rows = list(reader)
    if not rows:
        return "[Empty CSV file]"
    lines = ["Columns: " + ", ".join(rows[0]), "Data Rows:"]
    for r in rows[1:150]:
        lines.append(" | ".join(r))
    return "\n".join(lines)

def extract_xlsx_text(file_bytes: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        lines.append(f"--- Sheet: {sheet_name} ---")
        sheet = wb[sheet_name]
        for row in sheet.iter_rows(values_only=True):
            row_vals = [str(v) for v in row if v is not None]
            if row_vals:
                lines.append(" | ".join(row_vals))
    return "\n".join(lines) if lines else "[Empty XLSX workbook]"

def extract_file_content(file_bytes: bytes, filename: str, mime_type: str) -> str:
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf') or 'pdf' in mime_type.lower():
        return extract_pdf_text(file_bytes)
    elif filename_lower.endswith('.docx') or 'wordprocessingml' in mime_type.lower():
        return extract_docx_text(file_bytes)
    elif filename_lower.endswith('.csv') or 'csv' in mime_type.lower():
        return extract_csv_text(file_bytes)
    elif filename_lower.endswith('.xlsx') or 'spreadsheetml' in mime_type.lower():
        return extract_xlsx_text(file_bytes)
    else:
        return extract_txt_text(file_bytes)

def normalize_attachment_context(filename: str, mime_type: str, raw_text: str) -> str:
    MAX_CHARS = 30000
    clean_text = raw_text.strip()
    truncated = len(clean_text) > MAX_CHARS
    if truncated:
        clean_text = clean_text[:MAX_CHARS]

    context = f"[ATTACHMENT]\nFilename: {filename}\nType: {mime_type}\n\n[CONTENT]\n{clean_text}\n"
    if truncated:
        context += "\n[File content truncated because it exceeds context limit.]\n"
    context += "[END ATTACHMENT]"
    return context

def call_hermes_cli(prompt: str) -> str:
    try:
        cmd = [HERMES_EXE, "chat", "-q", prompt, "-Q"]
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120
        )

        if result.returncode != 0:
            error_msg = result.stderr.strip() or f"Hermes process exited with code {result.returncode}"
            raise Exception(error_msg)

        output = result.stdout.strip()
        lines = output.splitlines()
        clean_lines = [line for line in lines if not line.startswith("session_id:")]
        final_text = "\n".join(clean_lines).strip()
        return final_text if final_text else "No response from AI."

    except subprocess.TimeoutExpired:
        raise Exception("Hermes CLI request timed out (120s).")
    except Exception as e:
        raise Exception(f"Failed to execute Hermes CLI: {e}")

def analyze_furniture_visualization_9router(image_url: str, user_prompt: str, current_design_state: dict) -> str:
    """
    Calls 9Router Vision API (gemini/gemini-3.7-flash) to inspect furniture visualization objectively.
    Returns structured vision observation string or empty string on failure/timeout.
    """
    try:
        nine_router_base = os.environ.get('NINE_ROUTER_BASE_URL', 'http://localhost:20128/v1').rstrip('/')
        nine_router_key = os.environ.get('NINE_ROUTER_API_KEY', '')
        
        system_instruction = (
            "You are the visual inspection component of AGM Assistant. "
            "Analyze the provided AGM furniture visualization objectively. "
            "Do NOT infer exact physical dimensions from the image. "
            "Do NOT invent technical specifications. "
            "Focus only on visible characteristics: furniture structure, leg structure, proportions, color, material appearance, finish, and style. "
            f"Compare against state: {json.dumps(current_design_state or {})}. "
            f"Customer feedback: '{user_prompt}'. "
            "Return concise JSON with keys: 'observations', 'feedback_target', 'affected_visual_area', 'confidence'."
        )

        payload = {
            "model": "gemini/gemini-3.7-flash",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": system_instruction},
                        {"type": "image_url", "image_url": {"url": image_url}}
                    ]
                }
            ]
        }

        data_bytes = json.dumps(payload).encode('utf-8')
        req = urllib.request.Request(
            f"{nine_router_base}/chat/completions",
            data=data_bytes,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {nine_router_key}" if nine_router_key else ""
            }
        )

        with urllib.request.urlopen(req, timeout=15) as resp:
            resp_bytes = resp.read()
            raw_text = resp_bytes.decode('utf-8')
            
            # Handle streaming chunks if returned as SSE
            if "data: " in raw_text:
                chunks = raw_text.split("data: ")
                extracted_content = ""
                for c in chunks:
                    c_clean = c.strip()
                    if c_clean and c_clean != "[DONE]":
                        try:
                            parsed_c = json.loads(c_clean)
                            delta = parsed_c.get('choices', [{}])[0].get('delta', {}).get('content', '')
                            extracted_content += delta
                        except Exception:
                            pass
                return extracted_content.strip()
            else:
                parsed_res = json.loads(raw_text)
                return parsed_res.get('choices', [{}])[0].get('message', {}).get('content', '').strip()

    except Exception as vision_err:
        print(f"[{time.strftime('%H:%M:%S')}] Warning: 9Router Vision analysis skipped/failed: {vision_err}")
        return ""

def process_ai_job(job):
    job_id = job['id']
    conversation_id = job['conversation_id']
    raw_message = job['message']
    incoming_design_state = job.get('design_state')
    
    user_prompt = raw_message
    attachment_info = None

    try:
        if raw_message.startswith('{'):
            parsed = json.loads(raw_message)
            if isinstance(parsed, dict):
                if 'text' in parsed:
                    user_prompt = parsed['text']
                if 'attachment' in parsed:
                    attachment_info = parsed['attachment']
                if 'currentDesignState' in parsed and parsed['currentDesignState']:
                    incoming_design_state = parsed['currentDesignState']
    except Exception as parse_err:
        pass

    print(f"[{time.strftime('%H:%M:%S')}] Processing job {job_id} for session {conversation_id}: '{user_prompt[:50]}...'")

    try:
        supabase.from_("ai_jobs").update({
            "status": "processing",
            "processing_start_at": "now()"
        }).eq("id", job_id).execute()

        attachment_prefix = ""
        if attachment_info and attachment_info.get('storage_url'):
            url = attachment_info['storage_url']
            filename = attachment_info.get('filename', 'attachment')
            mime_type = attachment_info.get('mime_type', 'application/octet-stream')
            source_type = attachment_info.get('source', 'furniture_reference')
            
            is_image_attachment = any(mime_type.lower().startswith(p) for p in ['image/jpeg', 'image/png', 'image/webp']) or any(filename.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.webp'])

            if is_image_attachment:
                print(f"[{time.strftime('%H:%M:%S')}] Customer image attachment ({source_type}) detected. Calling 9Router Vision...")
                header_tag = "[ROOM VISUAL ANALYSIS]" if source_type == 'room_photo' else "[FURNITURE REFERENCE ANALYSIS]" if source_type == 'furniture_reference' else "[DESIGN INSPIRATION ANALYSIS]"
                img_vision_res = analyze_furniture_visualization_9router(url, f"Analyze this customer uploaded image ({source_type}). {user_prompt}", incoming_design_state)
                if img_vision_res:
                    attachment_prefix = f"{header_tag}\nFilename: {filename}\nSource: {source_type}\nVisual Analysis: {img_vision_res}\n[END ATTACHMENT]\n\n"
                else:
                    attachment_prefix = f"{header_tag}\nFilename: {filename}\nSource: {source_type}\n[Image URL: {url}]\n[END ATTACHMENT]\n\n"
            else:
                req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req) as resp:
                    file_bytes = resp.read()

                extracted_text = extract_file_content(file_bytes, filename, mime_type)
                attachment_prefix = normalize_attachment_context(filename, mime_type, extracted_text) + "\n\n"

        # Vision Pipeline Execution Guard (Trigger Vision ONLY if active visualization image exists AND prompt implies visual feedback)
        vision_context = ""
        active_img_url = (incoming_design_state or {}).get('visualization', {}).get('imageUrl')
        
        # Check if prompt contains visual feedback indicators (leg, color, proportion, shape, appearance, etc.)
        prompt_lower = user_prompt.lower()
        is_explicit_dim = any(k in prompt_lower for k in ['panjang', 'lebar', 'tinggi', 'kedalaman', 'cm', 'mm', 'ukuran'])
        visual_keywords = ['kaki', 'warna', 'proporsi', 'bentuk', 'terlalu besar', 'terlalu kecil', 'terlalu terang', 'terlalu gelap', 'desainnya', 'kelihatan', 'tampak', 'elegan']
        is_visual_feedback = any(vk in prompt_lower for vk in visual_keywords) and not is_explicit_dim

        if active_img_url and is_visual_feedback:
            print(f"[{time.strftime('%H:%M:%S')}] Active visualization image detected. Triggering 9Router Vision analysis...")
            vision_result = analyze_furniture_visualization_9router(active_img_url, user_prompt, incoming_design_state)
            if vision_result:
                vision_context = f"[9ROUTER VISION MODEL INSPECTION RESULT]\n{vision_result}\n(Gunakan hasil analisis visual ini sebagai panduan konteks visual produk untuk menjawab umpan balik customer).\n\n"

        # Construct Previous State Prompt Context
        state_context = ""
        if incoming_design_state and isinstance(incoming_design_state, dict) and incoming_design_state.get('category'):
            state_context = f"[CURRENT DESIGN STATE IN SESSION]\n```json\n{json.dumps(incoming_design_state, indent=2)}\n```\n(Gunakan state ini sebagai rujukan utama. Jika user melakukan perubahan kecil, pertahankan seluruh field lain dan naikkan version +1).\n\n"
        else:
            state_context = "[CURRENT DESIGN STATE IN SESSION]\n(Belum ada draf desain aktif. Jangan mengarang desain kecuali user meminta membuat custom furniture).\n\n"

        # Combine System Persona + State Context + Vision Context + Attachment + User Prompt
        full_prompt = f"{SYSTEM_CONSULTANT_INSTRUCTION}\n\n{state_context}{vision_context}{attachment_prefix}Pertanyaan/Instruksi Customer:\n{user_prompt}"

        ai_response_text = call_hermes_cli(full_prompt)

        # Parse & Validate Output Design State
        updated_design_state = incoming_design_state
        try:
            if "```json_design_state" in ai_response_text:
                parts = ai_response_text.split("```json_design_state")
                if len(parts) > 1:
                    json_str = parts[1].split("```")[0].strip()
                    parsed = json.loads(json_str)
                    if isinstance(parsed, dict) and parsed.get('category'):
                        # --- CANONICAL STATE NORMALIZATION ---
                        # 1. Category Normalization
                        raw_cat = str(parsed.get('category', '')).lower()
                        if 'meja makan' in raw_cat or 'dining' in raw_cat:
                            parsed['category'] = 'dining_table'
                        elif 'lemari' in raw_cat or 'wardrobe' in raw_cat:
                            parsed['category'] = 'wardrobe'
                        elif 'sofa' in raw_cat:
                            parsed['category'] = 'sofa'
                        elif 'tv' in raw_cat or 'credenza' in raw_cat:
                            parsed['category'] = 'tv_cabinet'
                        elif 'kitchen' in raw_cat:
                            parsed['category'] = 'kitchen_set'
                        elif 'kursi' in raw_cat or 'chair' in raw_cat:
                            parsed['category'] = 'chair'
                        elif 'meja' in raw_cat or 'table' in raw_cat:
                            parsed['category'] = 'table'
                        else:
                            parsed['category'] = 'other'

                        # 2. Capacity Normalization to Integer
                        raw_cap = parsed.get('capacity')
                        if raw_cap is not None:
                            import re
                            nums = re.findall(r'\d+', str(raw_cap))
                            if nums:
                                parsed['capacity'] = int(nums[0])

                        # 3. Dimension Semantics Normalization (length/width/depth/height)
                        dims = parsed.get('dimensions')
                        if not isinstance(dims, dict):
                            dims = {}
                        
                        # Map legacy width if length is missing and width was used as length
                        legacy_width = parsed.get('width')
                        legacy_depth = parsed.get('depth')
                        legacy_height = parsed.get('height')
                        
                        if legacy_width and not dims.get('length'):
                            dims['length'] = int(legacy_width) if str(legacy_width).isdigit() else legacy_width
                        if legacy_depth and not dims.get('width'):
                            dims['width'] = int(legacy_depth) if str(legacy_depth).isdigit() else legacy_depth
                        if legacy_height and not dims.get('height'):
                            dims['height'] = int(legacy_height) if str(legacy_height).isdigit() else legacy_height

                        if 'unit' not in dims:
                            dims['unit'] = 'cm'

                        parsed['dimensions'] = dims
                        
                        # Remove legacy top-level dimension keys to enforce clean canonical schema
                        parsed.pop('width', None)
                        parsed.pop('depth', None)
                        parsed.pop('height', None)

                        # Versioning & Visualization State Lifecycle Guarantee
                        if incoming_design_state and isinstance(incoming_design_state, dict):
                            prev_ver = incoming_design_state.get('version', 1)
                            parsed['version'] = prev_ver + 1
                            
                            # Inherit or set stale visualization status if state changed
                            prev_vis = incoming_design_state.get('visualization', {})
                            new_vis = parsed.get('visualization', {})
                            
                            if prev_vis.get('status') == 'ready':
                                new_vis['status'] = 'stale'
                                new_vis['imageUrl'] = prev_vis.get('imageUrl') # Retain as previous reference
                                new_vis['designVersion'] = prev_vis.get('designVersion', prev_ver)
                                # Preserve history
                                hist = prev_vis.get('visualizationHistory', [])
                                if prev_vis.get('imageUrl') and not any(h.get('imageUrl') == prev_vis.get('imageUrl') for h in hist):
                                    hist.append({
                                        'designVersion': prev_vis.get('designVersion', prev_ver),
                                        'imageUrl': prev_vis.get('imageUrl'),
                                        'generatedAt': prev_vis.get('generatedAt')
                                    })
                                new_vis['visualizationHistory'] = hist
                            elif not new_vis.get('status'):
                                new_vis['status'] = 'not_configured'
                                new_vis['designVersion'] = parsed['version']
                                new_vis['visualizationHistory'] = prev_vis.get('visualizationHistory', [])
                            parsed['visualization'] = new_vis
                        elif not parsed.get('version'):
                            parsed['version'] = 1
                            if not parsed.get('visualization'):
                                parsed['visualization'] = {
                                    'status': 'not_configured',
                                    'designVersion': 1
                                }

                        updated_design_state = parsed
        except Exception as parse_state_err:
            print(f"[{time.strftime('%H:%M:%S')}] Warning: Failed parsing updated design_state: {parse_state_err}. Retaining previous state.")

        update_payload = {
            "status": "completed",
            "response": ai_response_text,
            "completed_at": "now()"
        }
        if updated_design_state is not None:
            try:
                update_payload["design_state"] = updated_design_state
            except Exception:
                pass

        try:
            supabase.from_("ai_jobs").update(update_payload).eq("id", job_id).execute()
        except Exception as update_err:
            if "design_state" in str(update_err):
                # Fallback if DB table has not cached design_state column yet
                update_payload.pop("design_state", None)
                supabase.from_("ai_jobs").update(update_payload).eq("id", job_id).execute()
            else:
                raise update_err

        print(f"[{time.strftime('%H:%M:%S')}] ✓ Job {job_id} completed successfully.")

    except Exception as e:
        error_message = str(e)
        print(f"[{time.strftime('%H:%M:%S')}] ✗ Job {job_id} failed: {error_message}")
        supabase.from_("ai_jobs").update({
            "status": "failed",
            "error": error_message,
            "completed_at": "now()"
        }).eq("id", job_id).execute()

def listen_for_new_jobs():
    print("==================================================")
    print(" AGM Assistant Furniture Consultant Worker Active")
    print(f" Target Supabase : {SUPABASE_URL}")
    print("==================================================")

    while True:
        try:
            res = supabase.from_("ai_jobs").select("*").eq("status", "pending").order("created_at", desc=False).limit(1).execute()
            jobs = res.data
            if jobs:
                for job in jobs:
                    process_ai_job(job)
        except Exception as e:
            print(f"[{time.strftime('%H:%M:%S')}] Error fetching pending jobs: {e}")
        time.sleep(3)

if __name__ == "__main__":
    listen_for_new_jobs()
